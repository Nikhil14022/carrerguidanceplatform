import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("HOLISTREE CAREER GUIDANCE PLATFORM")
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("Website Skeleton Diagram & System Architecture Document")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Metadata Box Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    meta_data = [
        [("Project:", " Career Guidance & Alignment Platform"), ("Document Version:", " 2.0 (Production Blueprint)")],
        [("Architecture:", " Multi-Portal Next.js App Router & Groq AI"), ("Date:", " August 2026")]
    ]
    
    for r_idx, row in enumerate(meta_table.rows):
        for c_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            k, v = meta_data[r_idx][c_idx]
            run_k = p.add_run(k)
            run_k.font.bold = True
            run_k.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run_v = p.add_run(v)
            run_v.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 1: Executive Overview
    h1 = doc.add_heading("1. Executive Overview & Multi-Portal Architecture", level=1)
    h1.style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p = doc.add_paragraph(
        "The Holistree Career Guidance Platform is structured as a role-governed, multi-portal web application designed to guide students through a structured 9-stage career alignment journey. The architecture connects four distinct user personas into a unified real-time ecosystem powered by Next.js 16, Prisma ORM, MongoDB, and Groq AI."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)

    # Role Overview Table
    role_table = doc.add_table(rows=5, cols=4)
    role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Portal Role", "Base Route Path", "Primary Responsibilities", "Key System Features"]
    hdr_cells = role_table.rows[0].cells
    for idx, heading in enumerate(headers):
        set_cell_background(hdr_cells[idx], "0F172A")
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(heading)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    roles_info = [
        ("Student / Client", "/dashboard/*", "Complete 14 assessment modules, explore AI research, track 9 journey stages.", "Interactive Form Editor, AI Research Lab, Reports Feed, Expert Chat"),
        ("Parent Account", "/parent/*", "Monitor child's assessment completion, review recommendations & progress.", "Linked Student Dashboard, Stage Progress Tracker, Report View"),
        ("Mentor / Counselor", "/mentor/*", "Review module answers, manage 9-stage progression, generate AI reports.", "Combined Answers Feed, Decluttered Stage Tracker, Live Chat, Report Editor"),
        ("Super Admin", "/admin/*", "Supervise platform users, assign mentors to clients, manage system modules.", "Tabbed Mentors & Clients Roster, System Config, Audit & User Access Control")
    ]

    for r_idx, (r_name, r_route, r_resp, r_feat) in enumerate(roles_info, start=1):
        row_cells = role_table.rows[r_idx].cells
        bg_color = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        
        for c_idx, val in enumerate([r_name, r_route, r_resp, r_feat]):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif c_idx == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 2: Website Skeleton & Route Hierarchy Diagram
    h2 = doc.add_heading("2. Website Skeleton & Navigation Map", level=1)
    h2.style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph("The following diagram demarkates the complete structural tree and route hierarchy across public, student, parent, mentor, and administrative routes:")

    # Skeleton Box / Text Diagram
    diagram_p = doc.add_paragraph()
    diagram_p.paragraph_format.line_spacing = 1.0
    diagram_p.paragraph_format.space_after = Pt(12)
    
    skeleton_ascii = """
[ PLATFORM ROOT: / ]
 │
 ├── 🔑 AUTHENTICATION GATEWAY
 │    ├── /login                     (Multi-role credential login & NextAuth session generation)
 │    └── /register                  (Client & Linked Parent registration form)
 │
 ├── 🎓 STUDENT PORTAL [/dashboard]
 │    ├── /dashboard                 (Profile Overview & 9-Stage Completion Stepper)
 │    ├── /dashboard/modules         (14 Assessment Modules Grid: Self-Discovery, Skills, Academic)
 │    │    └── /dashboard/modules/[id]  (Interactive Question-by-Question Form & Scale Editor)
 │    ├── /dashboard/research        (AI Profession Research Lab: 2-Stage Intelligence Engine)
 │    ├── /dashboard/reports         (Generated Career Alignment Reports & Final Plans)
 │    ├── /dashboard/resume          (Structured Student Resume & Portfolio Builder)
 │    ├── /dashboard/expert-chat     (Real-time Counselor & AI Support Messaging)
 │    └── /dashboard/appointments    (Google Calendar Synced Counseling Sessions)
 │
 ├── 👪 PARENT PORTAL [/parent]
 │    ├── /parent                    (Parent Dashboard & Child Milestone Overview)
 │    ├── /parent/progress           (Child Stage Progress & Assessment Completion Metrics)
 │    ├── /parent/client-data        (View-Only Access to Child Submitted Answers)
 │    └── /parent/questionnaire      (Parent Insights & Family Expectation Assessment)
 │
 ├── 👨‍🏫 MENTOR PORTAL [/mentor]
 │    ├── /mentor                    (Mentor Dashboard & Assigned Client Roster)
 │    ├── /mentor/clients/[id]       (Client Journey Control Panel & Decluttered Stage Stepper)
 │    │    ├── /combined-answers     (Full Read-Only Feed of 14 Assessment Responses)
 │    │    └── /modules/[moduleId]   (Interactive Answer Editor for Structured Form Fields)
 │    └── /mentor/reports/[id]       (AI Draft Report Review & Counselor Human Finalisation)
 │
 └── 🛡️ SUPER ADMIN PORTAL [/admin]
      ├── /admin                     (Tabbed Roster: Mentors | Clients & Linked Parent Accounts)
      ├── /admin/clients/[id]        (Admin Client Deep Dive & Global Access Rights Control)
      ├── /admin/modules             (Global Module Schema & Question Order Configuration)
      ├── /admin/workflow            (Global Stage Definition & System Threshold Rules)
      └── /admin/analytics           (Platform Usage Statistics & Completion Metrics)
    """
    
    # Render ASCII Diagram inside shaded box table
    diag_table = doc.add_table(rows=1, cols=1)
    diag_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    diag_cell = diag_table.rows[0].cells[0]
    set_cell_background(diag_cell, "0F172A")
    set_cell_margins(diag_cell, top=140, bottom=140, left=180, right=180)
    
    p_code = diag_cell.paragraphs[0]
    p_code.paragraph_format.space_after = Pt(0)
    run_code = p_code.add_run(skeleton_ascii.strip())
    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(8.5)
    run_code.font.color.rgb = RGBColor(0xFB, 0xBF, 0x24)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 3: 9-Stage Career Progression Lifecycle Skeleton
    h3 = doc.add_heading("3. Student 9-Stage Career Journey Lifecycle", level=1)
    h3.style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph("The core methodology of the platform is organized into 9 sequential progression stages. Each stage is tracked visually using a decluttered horizontal timeline stepper:")

    stages_data = [
        ("Stage 1: Self-Discovery & Personality Assessment", "Modules 1-3", "Analyze core personality traits, interests, learning style, and personal values."),
        ("Stage 2: Aptitude & Skill Mapping", "Modules 4-5", "Evaluate cognitive capabilities, logical reasoning, numerical and verbal strengths."),
        ("Stage 3: Domain & Career Exploration", "AI Research Lab", "Examine profession profiles, daily tasks, lifestyle trade-offs, and industry demand."),
        ("Stage 4: Academic & Stream Alignment", "Modules 6-8", "Select secondary school streams (Science/Commerce/Arts) and undergraduate degree options."),
        ("Stage 5: College & Entrance Exam Strategy", "Modules 9-10", "Identify target Indian & foreign universities, qualifying exams (JEE, NEET, CUET, SAT)."),
        ("Stage 6: Profile & Skill Gap Identification", "Modules 11-12", "Benchmark student capabilities against target college entry criteria and industry skills."),
        ("Stage 7: Portfolio, Projects & Resume Building", "Resume Builder", "Document extracurricular achievements, certifications, internships, and project work."),
        ("Stage 8: Practical Exposure & Interview Readiness", "Modules 13-14", "Conduct mock counselor interviews, shadow professionals, and participate in workshops."),
        ("Stage 9: Final Career Action Plan & Execution", "Final Report", "Issue finalized 12-month execution roadmap co-signed by counselor, parent, and student.")
    ]

    stage_table = doc.add_table(rows=10, cols=3)
    stage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    stg_hdr = stage_table.rows[0].cells
    for idx, text in enumerate(["Stage Name & Title", "Associated Module / Tool", "Stage Objective & Deliverables"]):
        set_cell_background(stg_hdr[idx], "0F172A")
        p = stg_hdr[idx].paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for r_idx, (s_name, s_tool, s_obj) in enumerate(stages_data, start=1):
        row_cells = stage_table.rows[r_idx].cells
        bg_color = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        
        for c_idx, val in enumerate([s_name, s_tool, s_obj]):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif c_idx == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 4: Data Model & Database Blueprint
    h4 = doc.add_heading("4. Database Schema & Entity Relationships Diagram", level=1)
    h4.style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph("The underlying MongoDB database managed via Prisma ORM establishes strict entity relations between users, profiles, assessment modules, stage milestones, and reports:")

    db_ascii = """
┌──────────────────────────────────┐        ┌──────────────────────────────────┐
│              User                │        │          ClientProfile           │
├──────────────────────────────────┤        ├──────────────────────────────────┤
│ id           String (ObjectId)   │◄───────┤ id           String (ObjectId)   │
│ email        String (Unique)     │        │ userId       String (Ref: User) │
│ password     String (Hashed)     │        │ parentId     String (Ref: User) │
│ role         Enum (Role)         │        │ currentStage Int (1-9)           │
│ name         String              │        │ journeyStatus String             │
└──────────────────────────────────┘        └──────────────────────────────────┘
                 ▲                                            ▲
                 │ (1:1)                                      │ (1:N)
┌────────────────┴─────────────────┐        ┌─────────────────┴────────────────┐
│          MentorProfile           │        │           ClientModule           │
├──────────────────────────────────┤        ├──────────────────────────────────┤
│ id           String (ObjectId)   │        │ id           String (ObjectId)   │
│ userId       String (Ref: User) │        │ clientProfileId Ref(Client)       │
│ type         Enum (PERMANENT)    │        │ moduleId     String (e.g. mod1)  │
│ status       Enum (ACTIVE)       │        │ status       Enum (SUBMITTED...) │
└──────────────────────────────────┘        │ responseData Json (Form Answers) │
                 ▲                          └──────────────────────────────────┘
                 │ (1:N)                                      ▲
┌────────────────┴─────────────────┐                          │ (1:N)
│         MentorAssignment         │        ┌─────────────────┴────────────────┐
│ clientProfileId / mentorProfile  │        │           ClientStage            │
└──────────────────────────────────┘        │ stageNumber  Int (1-9)           │
                                            │ status       Enum (IN_PROGRESS) │
                                            └──────────────────────────────────┘
    """

    db_table = doc.add_table(rows=1, cols=1)
    db_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    db_cell = db_table.rows[0].cells[0]
    set_cell_background(db_cell, "0F172A")
    set_cell_margins(db_cell, top=140, bottom=140, left=180, right=180)
    
    p_db = db_cell.paragraphs[0]
    p_db.paragraph_format.space_after = Pt(0)
    run_db = p_db.add_run(db_ascii.strip())
    run_db.font.name = 'Consolas'
    run_db.font.size = Pt(8.5)
    run_db.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 5: AI Research & Intelligence Pipeline
    h5 = doc.add_heading("5. AI Research & Intelligence Pipeline Skeleton", level=1)
    h5.style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "The platform embeds Groq AI (llama-3.1-8b-instant) to dynamically synthesize structured profession intelligence and student profile alignment reports:"
    )

    ai_steps = [
        ("1. Query Reception & Prompt Construction", "The student submits a profession title (e.g. 'Aerospace Engineer'). The system fetches the student's completed 14 module assessment responses from MongoDB."),
        ("2. Groq AI Inference & Structured JSON Enforcement", "A system prompt enforces strict JSON schema generation with strict token budget limits (max 3500 tokens) to guarantee rapid response times."),
        ("3. Two-Stage Intel Synthesis", "Stage 1 builds 14 basic understanding cards (language, daily duties, entrance exams). Stage 2 builds advanced specialisations, salary timelines, and fit scores."),
        ("4. Safe Content Formatter & UI Rendering", "The frontend processes raw JSON into formatted cards, timeline grids, and diagnostic pills using explicit high-contrast inline styles.")
    ]

    for step_title, step_desc in ai_steps:
        p_step = doc.add_paragraph()
        p_step.paragraph_format.left_indent = Inches(0.2)
        p_step.paragraph_format.space_after = Pt(4)
        run_st = p_step.add_run(f"• {step_title}: ")
        run_st.font.bold = True
        run_st.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        run_sd = p_step.add_run(step_desc)
        run_sd.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Footer note
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ft = footer_p.add_run("— End of Skeleton Architecture Document —")
    run_ft.font.italic = True
    run_ft.font.size = Pt(9.5)
    run_ft.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    file_path = "Website_Skeleton_Diagram_Architecture.docx"
    doc.save(file_path)
    print(f"Document successfully created at {file_path}")

if __name__ == "__main__":
    create_document()
