import docx
import json

def extract_full_content(filepath, output_file):
    doc = docx.Document(filepath)
    lines = []
    
    # Extract all paragraphs
    lines.append("=== PARAGRAPHS ===\n")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            lines.append(f"P{i} [{p.style.name}]: {text}")
    
    # Extract all tables
    lines.append("\n\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(f"  Row {r_idx}: {' | '.join(cells)}")
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    print(f"Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables to {output_file}")

# Extract test document
extract_full_content(
    '2025 8-12th & Degree .docx',
    'extracted_test_content.txt'
)

# Extract interpretation document
extract_full_content(
    'Interpretation document for interns (Khushi) 2025.docx',
    'extracted_interpretation_content.txt'
)
